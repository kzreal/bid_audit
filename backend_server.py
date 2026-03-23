"""
HiAgent API 后端服务器
连接前端 Vue 应用和 HiAgent API
"""

from flask import Flask, request, jsonify
from flask_cors import CORS
import os
import sys
import logging
import tempfile
import base64
import requests
from pathlib import Path

# 配置日志
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 添加当前目录到 Python 路径，以便导入 hiagent_client
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from hiagent_client import TaskCreator, TaskAuditor, SummaryAgent
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

# ========== 千问 VL OCR 配置 ==========
QWEN_VL_API_KEY = os.getenv('QWEN_VL_API_KEY', 'sk-a1cd42b6-591e-4fb9-9af6-eb980060eb73').strip()
QWEN_VL_BASE_URL = os.getenv('QWEN_VL_BASE_URL', 'https://ai-model.chint.com/api').strip()
LLM_AVAILABLE = bool(QWEN_VL_API_KEY and QWEN_VL_BASE_URL)

print(f"[OCR] LLM Available: {LLM_AVAILABLE}")


class ImageRecognitionService:
    """图像识别服务 - 通过千问 VL API 识别图像内容"""

    def __init__(self, endpoint: str, api_key: str, timeout: int = 300, max_retries: int = 3):
        self.endpoint = endpoint
        self.api_key = api_key
        self.timeout = timeout
        self.max_retries = max_retries

        self.session = requests.Session()
        self.session.headers.update({
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        })

        retry_adapter = requests.adapters.HTTPAdapter(
            max_retries=max_retries,
            pool_connections=10,
            pool_maxsize=100
        )
        self.session.mount('http://', retry_adapter)
        self.session.mount('https://', retry_adapter)

    def describe_image(self, image_data: bytes, image_format: str) -> str:
        """
        发送图像到 LLM API 获取描述

        Args:
            image_data: 图像二进制数据
            image_format: 图像格式（如 'png', 'jpeg'）

        Returns:
            图像描述文本，失败返回 None
        """
        for attempt in range(self.max_retries):
            try:
                base64_image = base64.b64encode(image_data).decode('utf-8')

                request_data = {
                    'model': 'qwen-vl',
                    'messages': [
                        {
                            'role': 'user',
                            'content': [
                                {
                                    'type': 'text',
                                    'text': '十字以内直接概括图片的内容，不用加"图片展示..."、"这是...图片"等叙述'
                                },
                                {
                                    'type': 'image_url',
                                    'image_url': {
                                        'url': f'data:image/{image_format};base64,{base64_image}'
                                    }
                                }
                            ]
                        }
                    ],
                    'max_tokens': 100
                }

                response = self.session.post(
                    self.endpoint,
                    json=request_data,
                    timeout=self.timeout
                )

                response.raise_for_status()
                result = response.json()
                content = result['choices'][0]['message']['content'].strip()
                return content

            except Exception as e:
                logger.error(f"LLM API request failed (attempt {attempt + 1}/{self.max_retries}): {e}")
                if attempt == self.max_retries - 1:
                    return None
                import time
                wait_time = 2 ** attempt
                time.sleep(wait_time)

        return None

    def close(self):
        self.session.close()


# 初始化 OCR 服务
ocr_service = None
if LLM_AVAILABLE:
    ocr_service = ImageRecognitionService(
        endpoint=f"{QWEN_VL_BASE_URL}/chat/completions",
        api_key=QWEN_VL_API_KEY,
        timeout=300,
        max_retries=3
    )
    print(f"[OCR] OCR service initialized")

app = Flask(__name__)
CORS(app)  # 允许跨域请求

# 从环境变量或配置文件读取配置
def get_config():
    # 首先尝试从本地 .env 文件读取
    api_url = None
    task_creator_api_key = None
    task_auditor_api_key = None
    summary_api_key = None
    user_id = 'user001'

    try:
        with open('.env', 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line.startswith('VITE_API_BASE_URL='):
                    api_url = line.split('=', 1)[1]
                elif line.startswith('VITE_TASK_CREATOR_API_KEY='):
                    task_creator_api_key = line.split('=', 1)[1]
                elif line.startswith('VITE_TASK_AUDITOR_API_KEY='):
                    task_auditor_api_key = line.split('=', 1)[1]
                elif line.startswith('VITE_SUMMARY_API_KEY='):
                    summary_api_key = line.split('=', 1)[1]
                elif line.startswith('VITE_HIAGENT_USER_ID='):
                    user_id = line.split('=', 1)[1]
    except:
        pass

    # 如果本地文件没有设置，尝试从环境变量读取
    if not api_url:
        api_url = os.getenv('VITE_API_BASE_URL')
    if not task_creator_api_key:
        task_creator_api_key = os.getenv('VITE_TASK_CREATOR_API_KEY', 'd6ntpsf4piphvinbnmh0')
    if not task_auditor_api_key:
        task_auditor_api_key = os.getenv('VITE_TASK_AUDITOR_API_KEY', 'd6oo7pn4piphvinbrb6g')
    if not summary_api_key:
        summary_api_key = os.getenv('VITE_SUMMARY_API_KEY', 'd6oo7pn4piphvinbrb7g')
    if user_id == 'user001':
        user_id = os.getenv('VITE_HIAGENT_USER_ID', '250701283')

    # 如果都失败，使用默认值
    if not api_url:
        api_url = 'https://prd-ai-studio.chint.com/api/proxy/api/v1'

    return api_url, task_creator_api_key, task_auditor_api_key, summary_api_key, user_id

API_URL, TASK_CREATOR_API_KEY, TASK_AUDITOR_API_KEY, SUMMARY_API_KEY, USER_ID = get_config()
print(f"Config: API_URL={API_URL}, TASK_CREATOR_API_KEY={TASK_CREATOR_API_KEY}, TASK_AUDITOR_API_KEY={TASK_AUDITOR_API_KEY}, SUMMARY_API_KEY={SUMMARY_API_KEY}, USER_ID={USER_ID}")

# 初始化 HiAgent 客户端，使用不同的 API key
task_creator = TaskCreator(API_URL, TASK_CREATOR_API_KEY, USER_ID)
task_auditor = TaskAuditor(API_URL, TASK_AUDITOR_API_KEY, USER_ID)
summary_agent = SummaryAgent(API_URL, SUMMARY_API_KEY, USER_ID)


@app.route('/hiagent/generate-tasks', methods=['POST'])
def generate_tasks():
    """生成审核任务"""
    try:
        data = request.get_json()
        requirement = data.get('requirement')
        # 检查是否明确传递了 type 参数
        task_type = data.get('type') if 'type' in data else None

        if not requirement:
            return jsonify({
                'code': 400,
                'message': '招标文件信息不能为空'
            }), 400

        # 所有需求都直接调用 HiAgent LLM 生成任务
        # 如果传递了 type 参数，使用传递的值；否则使用默认值 1
        type_param = task_type if task_type is not None else 1

        result = task_creator.sync_run_workflow({
            "extraction": requirement,
            "type": type_param
        })

        if not result or result.get('status') != 'success':
            return jsonify({
                'code': 500,
                'message': '任务生成失败'
            }), 500

        # 获取 output 字段（HiAgent API 新格式）
        output = result.get('output', '')

        # 将 output 传递给 TaskCreator.parse_tasks 进行解析
        # 新格式: {"output": "{\"tasks\": [\"任务1\", \"任务2\"]}"}
        tasks = TaskCreator.parse_tasks(output)

        return jsonify({
            'code': 200,
            'message': '任务生成成功',
            'data': tasks,
            'raw_text': output
        })

    except Exception as e:
        import traceback
        error_msg = f"生成任务时发生错误：{str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return jsonify({
            'code': 500,
            'message': '任务生成失败',
            'error': str(e),
            'traceback': traceback.format_exc()
        }), 500


@app.route('/hiagent/generate-conclusion', methods=['POST'])
def generate_conclusion():
    """生成最终审核结论"""
    try:
        data = request.get_json()
        task_input = data.get('task')
        reviews = data.get('reviews', [])

        if not task_input:
            return jsonify({
                'code': 400,
                'message': '任务不能为空'
            }), 400

        # 如果 task 是对象，从中提取任务描述 - 优先使用 title
        if isinstance(task_input, dict):
            # 优先使用 title，如果不存在再使用 description
            task = task_input.get('title', task_input.get('description', ''))
            # 如果 title 也是空，使用 description 作为 fallback
            if not task:
                task = task_input.get('description', '')
        else:
            # 如果是字符串，直接使用
            task = str(task_input)

        if not isinstance(reviews, list):
            return jsonify({
                'code': 400,
                'message': '审核结果必须是数组格式'
            }), 400

        # 调用 SummaryAgent 生成结论
        result_text = summary_agent.generate_conclusion(task, reviews, use_sync=True)

        if not result_text:
            return jsonify({
                'code': 500,
                'message': '结论生成失败'
            }), 500

        # 使用 SummaryAgent.parse_conclusion 解析总结结果
        parsed = summary_agent.parse_conclusion(result_text)
        conclusion = parsed.get('conclusion', '')
        reason = parsed.get('reason', '')
        evidence = parsed.get('evidence', '')

        # 调试：打印解析结果
        print(f"\n=== parse_conclusion 结果 ===")
        print(f"conclusion: '{conclusion}'")
        print(f"reason: '{reason}'")
        print(f"evidence: '{evidence}'")

        # 直接使用 hiagent 输出的 conclusion 作为 status
        status = conclusion if conclusion else '待确认'

        print(f"最终 status: '{status}'")

        # 使用 Python 的 datetime 获取当前时间
        from datetime import datetime

        # 返回符合 guide.md 定义的格式
        return jsonify({
            'code': 200,
            'message': '总结成功',
            'data': {
                'conclusion': conclusion,
                'reason': reason,
                'evidence': evidence
            },
            'status': status,
            'raw_text': result_text
        })

    except Exception as e:
        print(f"生成结论时发生错误：{str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误：{str(e)}'
        }), 500


@app.route('/hiagent/review-task', methods=['POST'])
def review_task():
    """审核任务"""
    try:
        data = request.get_json()
        task_input = data.get('task')
        context = data.get('context')

        if not task_input:
            return jsonify({
                'code': 400,
                'message': '任务不能为空'
            }), 400

        # 如果 task 是对象，从中提取任务描述 - 优先使用 title
        if isinstance(task_input, dict):
            # 优先使用 title，如果不存在再使用 description
            task = task_input.get('title', task_input.get('description', ''))
            # 如果 title 也是空，使用 description 作为 fallback
            if not task:
                task = task_input.get('description', '')
        else:
            # 如果是字符串，直接使用
            task = str(task_input)

        if not context:
            return jsonify({
                'code': 400,
                'message': '投标文件内容不能为空'
            }), 400

        # 调用 HiAgent API 审核任务
        result_text = task_auditor.audit_task(task, context, use_sync=True)

        if not result_text:
            return jsonify({
                'code': 500,
                'message': '任务审核失败'
            }), 500

        # 使用 TaskAuditor.parse_audit_result 解析审核结果
        # 新格式: {"result": {"suggestion": "...", "evidence": "..."}}
        parsed = TaskAuditor.parse_audit_result(result_text)
        suggestion = parsed.get('suggestion', '')
        evidence = parsed.get('evidence', '')

        # 返回符合格式
        return jsonify({
            'code': 200,
            'message': '任务审核成功',
            'data': {
                'suggestion': suggestion,
                'evidence': evidence
            },
            'raw_text': result_text
        })

    except Exception as e:
        print(f"审核任务时发生错误：{str(e)}")
        return jsonify({
            'code': 500,
            'message': f'服务器错误：{str(e)}'
        }), 500


@app.route('/hiagent/review-task-slices', methods=['POST'])
def review_task_slices():
    """
    多切片审核：对一个任务，用多个切片文件分别审核，然后汇总结果（不调用 LLM）
    """
    try:
        data = request.get_json()
        task_input = data.get('task')
        slices = data.get('slices', [])  # 切片内容数组

        # 验证切片数量（最多 100 个）
        if len(slices) > 100:
            return jsonify({'code': 400, 'message': '切片数量不能超过 100 个'}), 400

        if not task_input:
            return jsonify({'code': 400, 'message': '任务不能为空'}), 400

        if not slices:
            return jsonify({'code': 400, 'message': '切片不能为空'}), 400

        # 提取任务描述 - 优先使用 title，因为 description 可能为空字符串
        if isinstance(task_input, dict):
            # 优先使用 title，如果不存在再使用 description
            task = task_input.get('title', task_input.get('description', ''))
            # 如果 title 也是空，使用 description 作为 fallback
            if not task:
                task = task_input.get('description', '')
        else:
            task = str(task_input)

        print(f"开始多切片审核，任务：{str(task)[:50]}...，切片数：{len(slices)}")

        # 对每个切片调用 TaskAuditor
        reviews = []
        for idx, slice_text in enumerate(slices):
            print(f"正在审核切片 {idx+1}/{len(slices)}...")
            result_text = task_auditor.audit_task(task, slice_text, use_sync=True)
            parsed = TaskAuditor.parse_audit_result(result_text)

            # 处理 evidence，如果是字符串 "null" 则转为 null
            evidence = parsed.get('evidence', '')
            if evidence == 'null' or evidence == '""':
                evidence = None

            reviews.append({
                'suggestion': parsed.get('suggestion', ''),
                'evidence': evidence
            })

        # 调用 SummaryAgent 汇总所有切片的审核结果
        # 只返回原始切片审核结果，不做整合处理
        print(f"\n切片审核完成，共 {len(reviews)} 个切片")

        # 返回简化格式（只包含 task 和 reviews）
        return jsonify({
            'code': 200,
            'message': '多切片审核成功',
            'data': {
                'task': task,
                'reviews': reviews
            }
        })

    except Exception as e:
        import traceback
        error_msg = f"多切片审核失败：{str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return jsonify({
            'code': 500,
            'message': '多切片审核失败',
            'error': str(e)
        }), 500


@app.route('/hiagent/status', methods=['GET'])
def get_status():
    """获取 API 状态"""
    return jsonify({
        'code': 200,
        'message': 'API 服务正常',
        'data': {
            'api_url': API_URL,
            'user_id': USER_ID,
            'status': 'running'
        }
    })


@app.route('/health', methods=['GET'])
def health_check():
    """健康检查"""
    return jsonify({
        'status': 'ok',
        'message': 'HiAgent Backend Server is running'
    })


@app.route('/hiagent/summarize-reviews', methods=['POST'])
def summarize_reviews():
    """
    汇总切片审核结果：将各个切片的审核结果整合为统一格式
    """
    try:
        data = request.get_json()
        task_input = data.get('task')
        reviews_input = data.get('reviews', [])

        if not task_input:
            return jsonify({'code': 400, 'message': '任务不能为空'}), 400

        # 提取任务描述 - 优先使用 title，因为 description 可能为空字符串
        if isinstance(task_input, dict):
            # 优先使用 title，如果不存在再使用 description
            task = task_input.get('title', task_input.get('description', ''))
            # 如果 title 也是空，使用 description 作为 fallback
            if not task:
                task = task_input.get('description', '')
        else:
            task = str(task_input)

        # 检查 reviews 格式
        if not isinstance(reviews_input, list):
            return jsonify({'code': 400, 'message': 'reviews 必须是数组'}), 400

        print(f"\n开始汇总审核结果，任务：{str(task)[:50]}...，审核结果数：{len(reviews_input)}")

        # 构建 reviews 数组
        reviews = []
        for review in reviews_input:
            # 处理每个 review 对象
            suggestion = review.get('suggestion', '')
            evidence = review.get('evidence', '')

            # 如果 evidence 是字符串 "null"，则转为 null
            if evidence == 'null' or evidence == '""':
                evidence = None

            reviews.append({
                'suggestion': suggestion,
                'evidence': evidence
            })

        print(f"汇总完成，生成 reviews 数组：{len(reviews)} 条")

        # 返回结果
        return jsonify({
            'code': 200,
            'message': '汇总成功',
            'data': {
                'task': task,
                'reviews': reviews
            }
        })

    except Exception as e:
        import traceback
        error_msg = f"汇总失败：{str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return jsonify({
            'code': 500,
            'message': '汇总失败',
            'error': str(e)
        }), 500


# ========== 文档切片功能 ==========

def iter_block_items(parent):
    """遍历文档中的所有块元素（段落和表格），保持原始顺序"""
    if hasattr(parent, 'element'):
        parent_elm = parent.element.body
    else:
        parent_elm = parent

    for element in parent_elm.iterchildren():
        if isinstance(element, CT_P):
            yield Paragraph(element, parent)
        elif isinstance(element, CT_Tbl):
            yield Table(element, parent)


def get_heading_level(paragraph: Paragraph) -> int:
    """
    获取段落的标题级别
    Returns: 0=正文, 1=一级标题, 2=二级标题, ...
    """
    # 方法1: 检查样式名称
    if hasattr(paragraph, 'style') and paragraph.style.name:
        style_name = paragraph.style.name.lower()
        if 'heading 1' in style_name or '标题 1' in style_name:
            return 1
        elif 'heading 2' in style_name or '标题 2' in style_name:
            return 2
        elif 'heading 3' in style_name or '标题 3' in style_name:
            return 3
        elif 'heading 4' in style_name or '标题 4' in style_name:
            return 4
        elif 'heading 5' in style_name or '标题 5' in style_name:
            return 5

    # 方法2: 检查大纲级别
    if hasattr(paragraph, '_element'):
        p = paragraph._element
        if p.pPr is not None and hasattr(p.pPr, 'outlineLvl') and p.pPr.outlineLvl is not None:
            return int(p.pPr.outlineLvl.val) + 1

    return 0


def extract_paragraph_images(paragraph, doc) -> list:
    """
    检测段落中的图片

    Args:
        paragraph: 段落对象
        doc: 文档对象

    Returns:
        图片列表，每个图片包含 id, data, format
    """
    images = []

    # 方法1: 检查 wp:inline 格式的图片（内嵌）
    for run in paragraph.runs:
        for inline in run._element.xpath('.//w:drawing/wp:inline'):
            try:
                blip = inline.xpath('.//a:blip')
                if blip:
                    embed = inline.xpath('.//a:blip/@r:embed')
                    if embed:
                        image_data = _get_image_data(embed[0], doc)
                        if image_data:
                            image_format = _get_image_format(inline, doc)
                            images.append({
                                'id': embed[0],
                                'data': image_data,
                                'format': image_format
                            })
            except Exception as e:
                logger.warning(f"Error extracting wp:inline image: {e}")

    # 方法2: 检查 wp:anchor 格式的图片（锚定图片，如签名）
    for run in paragraph.runs:
        for drawing in run._element.xpath('.//w:drawing'):
            anchors = drawing.xpath('./wp:anchor')
            for anchor in anchors:
                try:
                    pics = anchor.xpath('.//pic:pic')
                    for pic in pics:
                        blips = pic.xpath('.//a:blip/@r:embed')
                        if blips:
                            rId = blips[0]
                            image_data = _get_image_data(rId, doc)
                            if image_data:
                                image_format = _get_image_format_from_rid(rId, doc)
                                images.append({
                                    'id': rId,
                                    'data': image_data,
                                    'format': image_format
                                })
                                logger.info(f"提取到 wp:anchor 图片 (ID: {rId}, 大小: {len(image_data)} 字节)")
                except Exception as e:
                    logger.warning(f"Error extracting wp:anchor image: {e}")

    return images


def _get_image_data(embed: str, doc) -> bytes:
    """从文档关系中提取图像数据"""
    try:
        if hasattr(doc, 'part') and hasattr(doc.part, 'related_parts'):
            image_part = doc.part.related_parts[embed]
            return image_part.blob
    except Exception:
        pass
    return None


def _get_image_format(image_element, doc) -> str:
    """检测图片格式"""
    try:
        blip = image_element.xpath('.//a:blip')
        if blip:
            embed = image_element.xpath('.//a:blip/@r:embed')
            if embed:
                image_part = doc.part.related_parts[embed[0]]
                return image_part.content_type.split('/')[-1]
    except Exception:
        pass
    return 'png'


def _get_image_format_from_rid(self, rId: str, doc) -> str:
    """从 relationship ID 获取图片格式"""
    try:
        image_part = doc.part.related_parts[rId]
        return image_part.content_type.split('/')[-1]
    except Exception:
        return 'png'


def process_images_batch(image_objects: list, ocr_svc) -> dict:
    """
    批量处理图片，使用 OCR 识别

    Args:
        image_objects: 图片对象列表
        ocr_svc: OCR 服务实例

    Returns:
        图片ID到描述的映射字典
    """
    processed_results = {}

    if not ocr_svc:
        logger.warning("OCR service not available, returning None for all images")
        return {img['id']: None for img in image_objects}

    for i, img in enumerate(image_objects, 1):
        try:
            logger.info(f"正在处理第 {i}/{len(image_objects)} 张图片: {img['id']}")
            image_size = len(img['data'])
            logger.info(f"图片大小: {image_size / 1024:.2f} KB, 格式: {img['format']}")

            if image_size > 10 * 1024 * 1024:  # 10MB
                processed_results[img['id']] = None
                logger.warning(f"Image {img['id']} too large, skipping")
                continue

            description = ocr_svc.describe_image(img['data'], img['format'])
            if description:
                processed_results[img['id']] = description
                logger.info(f"图片 {img['id']} 处理成功: {description}")
            else:
                processed_results[img['id']] = None
                logger.warning(f"图片 {img['id']} 识别失败")
        except Exception as e:
            logger.error(f"Error processing image {img['id']}: {e}")
            processed_results[img['id']] = None

    return processed_results


def slice_document(doc, max_level, ocr_svc=None):
    """按标题级别切片文档，支持图片 OCR 识别"""
    # 收集所有图片
    all_images = []
    logger.info(f"开始提取图片，文档包含 {len(doc.paragraphs)} 个段落和 {len(doc.tables)} 个表格")

    for paragraph in doc.paragraphs:
        images = extract_paragraph_images(paragraph, doc)
        all_images.extend(images)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    images = extract_paragraph_images(paragraph, doc)
                    all_images.extend(images)

    logger.info(f"共提取到 {len(all_images)} 张图片")

    # 使用 OCR 处理图片
    processed_images = {}
    if ocr_svc and all_images:
        logger.info(f"开始使用 OCR 处理 {len(all_images)} 张图片")
        processed_images = process_images_batch(all_images, ocr_svc)
        logger.info(f"OCR 图片处理完成，成功处理 {len([k for k, v in processed_images.items() if v])} 张图片")
    else:
        if not ocr_svc:
            logger.warning("OCR 服务不可用")
        for img in all_images:
            processed_images[img['id']] = None

    # 开始切片
    sections = []
    section_stack = []
    section_index = 0
    line_no = 1

    cover_section = {
        'level': 0,
        'title': '封面',
        'content': [],
        'index': section_index,
        'startLine': 1
    }
    section_stack.append(cover_section)

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            level = get_heading_level(block)
            text = block.text.strip()
            images = extract_paragraph_images(block, doc)

            if not text and not images:
                continue

            # 跳过目录
            if any(kw in text for kw in ['目录', '目  录', 'CONTENTS']):
                continue

            if level > 0 and level <= max_level:
                # 将当前栈顶章节添加到 sections，并从栈中移除
                if section_stack[-1]['content']:
                    section_stack[-1]['endLine'] = line_no - 1
                    sections.append(section_stack[-1])
                    section_index += 1
                    section_stack.pop()

                new_section = {
                    'level': level,
                    'title': text,
                    'content': [],
                    'index': section_index,
                    'startLine': line_no
                }

                while section_stack and section_stack[-1]['level'] > level:
                    section_stack.pop()

                section_stack.append(new_section)
                line_no += 1
            else:
                if level > 0:
                    section_stack[-1]['content'].append({'type': 'heading', 'level': level, 'text': text, 'line': line_no})
                elif text:
                    section_stack[-1]['content'].append({'type': 'paragraph', 'text': text, 'line': line_no})
                line_no += 1

            # 处理图片
            for img in images:
                description = processed_images.get(img['id'])
                if description:
                    section_stack[-1]['content'].append({'type': 'image', 'text': f'[图片: {description}]', 'line': line_no})
                else:
                    section_stack[-1]['content'].append({'type': 'image', 'text': '[图片: 未识别图片]', 'line': line_no})
                line_no += 1

        elif isinstance(block, Table):
            # 表格处理
            table_content = []
            for row in block.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip().replace('\n', ' '))
                table_content.append(row_data)

            if table_content:
                section_stack[-1]['content'].append({'type': 'table', 'data': table_content, 'line': line_no})
                line_no += len(table_content)

    # 添加栈中剩余的章节
    for section in section_stack:
        if section['content']:
            section['endLine'] = line_no - 1 if 'endLine' not in section else section['endLine']
            sections.append(section)

    return sections


@app.route('/document/slice', methods=['POST'])
def slice_document_endpoint():
    """
    切片文档接口
    接收 docx 文件和切片级别，返回切片列表
    """
    try:
        # 检查是否有文件
        if 'file' not in request.files:
            return jsonify({'code': 400, 'message': '没有上传文件'}), 400

        file = request.files['file']
        if file.filename == '':
            return jsonify({'code': 400, 'message': '文件名为空'}), 400

        # 获取切片级别 (0=整个文档, 1=一级标题, 2=二级标题, 3=三级标题)
        max_level = int(request.form.get('max_level', 0))

        print(f"开始切片文档: {file.filename}, max_level={max_level}")

        # 保存上传的文件到临时目录
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        try:
            # 使用 python-docx 加载文档
            doc = Document(tmp_path)

            # 创建 OCR 服务（如果可用）
            ocr_svc = None
            if LLM_AVAILABLE:
                try:
                    ocr_svc = ImageRecognitionService(
                        endpoint=f"{QWEN_VL_BASE_URL}/chat/completions",
                        api_key=QWEN_VL_API_KEY,
                        timeout=300,
                        max_retries=3
                    )
                    print("[OCR] 图像识别服务已创建")
                except Exception as e:
                    print(f"[OCR] 创建图像识别服务失败: {e}")
                    ocr_svc = None

            # 执行切片
            sections = slice_document(doc, max_level, ocr_svc)

            print(f"切片完成，共 {len(sections)} 个切片")

            # 返回切片信息
            return jsonify({
                'code': 200,
                'message': '切片成功',
                'data': {
                    'totalSlices': len(sections),
                    'slices': sections
                }
            })
        finally:
            # 删除临时文件
            Path(tmp_path).unlink(missing_ok=True)

    except Exception as e:
        import traceback
        error_msg = f"切片失败：{str(e)}\n{traceback.format_exc()}"
        print(error_msg)
        return jsonify({
            'code': 500,
            'message': '切片失败',
            'error': str(e)
        }), 500


@app.route('/document/slice/<int:slice_index>', methods=['GET'])
def get_slice_content(slice_index):
    """
    获取指定切片的内容
    需要先调用 /document/slice 上传文档
    """
    # 简化实现：直接返回切片信息，不缓存文档
    return jsonify({
        'code': 400,
        'message': '请先调用 /document/slice 上传文档'
    }), 400


# ========== 原有接口 ==========

if __name__ == '__main__':
    print("HiAgent Backend Server 启动中...")
    print(f"API URL: {API_URL}")
    print(f"User ID: {USER_ID}")
    print("服务器地址: http://localhost:8888")

    app.run(
        host='0.0.0.0',
        port=8888,
        debug=True,
        threaded=True
    )